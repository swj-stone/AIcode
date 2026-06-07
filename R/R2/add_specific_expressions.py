#!/usr/bin/env python3
"""Add Section 2.2 with specific fitted model expressions to the Word document."""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document(r'C:\Users\swj17\.claude\projects\R2\Practical_Work_2_Report.docx')

# Find insertion point: right before '3. Task 1' (Heading 1)
insert_idx = None
for i, p in enumerate(doc.paragraphs):
    if '3. Task 1' in p.text and 'Heading 1' in p.style.name:
        insert_idx = i
        break

if insert_idx is None:
    print('ERROR: Could not find Section 3')
    exit(1)

print(f'Inserting before paragraph {insert_idx}: "{doc.paragraphs[insert_idx].text[:80]}..."')

body = doc.element.body
target_element = body[insert_idx]
target_list_idx = list(body).index(target_element)

# Helpers to create elements detached from the document body
def make_heading(doc, text, level):
    p = doc.add_paragraph(text, style=f'Heading {level}')
    el = p._element
    body.remove(el)
    return el

def make_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    el = p._element
    body.remove(el)
    return el

def make_eq(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run('    ' + text)
    run.italic = True
    run.font.size = Pt(11)
    run.font.name = 'Cambria Math'
    el = p._element
    body.remove(el)
    return el

def make_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    p.add_run(text)
    el = p._element
    body.remove(el)
    return el

def make_empty(doc):
    p = doc.add_paragraph()
    el = p._element
    body.remove(el)
    return el

elements = []

# === Section 2.2 Header ===
elements.append(make_heading(doc, '2.2 Specific (Fitted) Model Expressions', 2))
elements.append(make_para(doc, 'The following are the fully specified fitted models with all estimated coefficients substituted into the equations. These are the concrete expressions that define each best model after parameter estimation on the full dataset.'))
elements.append(make_empty(doc))

# ======== TASK 1 ========
elements.append(make_heading(doc, '2.2.1 Task 1 -- Fitted Holt Linear Trend Model', 3))

elements.append(make_para(doc, 'Recurrence equations with optimised parameters (alpha = 0.13, beta = 0.99):', bold=True))
elements.append(make_eq(doc, 'L_t = 0.13 * Y_t + 0.87 * (L_{t-1} + T_{t-1})'))
elements.append(make_eq(doc, 'T_t = 0.99 * (L_t - L_{t-1}) + 0.01 * T_{t-1}'))
elements.append(make_empty(doc))

elements.append(make_para(doc, 'After fitting on all 84 quarterly observations (1960 Q1 to 1980 Q4), the final state is:', bold=True))
elements.append(make_eq(doc, 'L_84 = 15.1581   (level at 1980 Q4)'))
elements.append(make_eq(doc, 'T_84 =  0.0025   (trend at 1980 Q4)'))
elements.append(make_empty(doc))

elements.append(make_para(doc, 'h-step-ahead forecast from 1980 Q4 (for h = 1, 2, ..., 6 quarters):', bold=True))
elements.append(make_eq(doc, 'Y_hat_{84+h} = 15.1581 + 0.0025 * h'))
elements.append(make_empty(doc))

elements.append(make_para(doc, 'Interpretation: The near-zero trend (0.0025) means the forecast is essentially flat at approximately $15.16 per share. The high beta = 0.99 drives the trend almost entirely from the most recent level change, while the low alpha = 0.13 produces a heavily smoothed level estimate. MAPE on the 8-quarter control sample = 14.78%.'))
elements.append(make_empty(doc))

# ======== TASK 2 ========
elements.append(make_heading(doc, '2.2.2 Task 2 -- Fitted SARIMA(1,1,1)(0,1,1)[12] Model', 3))

elements.append(make_para(doc, 'Let Z_t = ln(Y_t) where Y_t is the monthly passenger count (thousands). Estimated coefficients:', bold=True))
elements.append(make_eq(doc, 'phi_1   (AR1)    = -0.2737'))
elements.append(make_eq(doc, 'theta_1 (MA1)    = -0.0878'))
elements.append(make_eq(doc, 'Theta_1 (SMA1)   = -0.5636'))
elements.append(make_eq(doc, 'sigma^2 (innovation variance) = 0.001362'))
elements.append(make_eq(doc, 'AIC = -477.96,  AICc = -477.48,  BIC = -463.59'))
elements.append(make_empty(doc))

elements.append(make_para(doc, 'Full SARIMA operator equation (substituting estimated coefficients):', bold=True))
elements.append(make_eq(doc, '(1 + 0.2737*B) (1 - B) (1 - B^12) Z_t  =  (1 - 0.0878*B) (1 - 0.5636*B^12) epsilon_t'))
elements.append(make_empty(doc))

elements.append(make_para(doc, 'Non-seasonal component: Let W_t = Z_t - Z_{t-1} be the first difference of log-passengers.', bold=True))
elements.append(make_eq(doc, 'W_t = -0.2737 * W_{t-1} + epsilon_t - 0.0878 * epsilon_{t-1}'))
elements.append(make_empty(doc))

elements.append(make_para(doc, 'Equivalently, expanded in terms of Z_t:', bold=True))
elements.append(make_eq(doc, 'Z_t = 0.7263 * Z_{t-1} + 0.2737 * Z_{t-2} + epsilon_t - 0.0878 * epsilon_{t-1}'))
elements.append(make_empty(doc))

elements.append(make_para(doc, 'Full SARIMA with seasonal component -- let W*_t = (1-B)(1-B^12)Z_t = Z_t - Z_{t-1} - Z_{t-12} + Z_{t-13}:', bold=True))
elements.append(make_eq(doc, 'W*_t = -0.2737 * W*_{t-1} + epsilon_t'))
elements.append(make_eq(doc, '       - 0.0878 * epsilon_{t-1} - 0.5636 * epsilon_{t-12} + 0.0495 * epsilon_{t-13}'))
elements.append(make_empty(doc))

elements.append(make_para(doc, 'Model diagnostics:', bold=True))
elements.append(make_bullet(doc, 'Ljung-Box (24 lags): p = 0.1839 -- residuals are white noise (no remaining autocorrelation structure).'))
elements.append(make_bullet(doc, 'Auto-ARIMA benchmark selected SARIMA(2,1,1)(0,1,1)[12] with AIC = -482.63 (marginally better, but both models yield nearly identical forecasts within 0.1% MAPE difference).'))
elements.append(make_empty(doc))

# ======== TASK 3 ========
elements.append(make_heading(doc, '2.2.3 Task 3 -- Fitted Multiplicative Classical Decomposition', 3))

elements.append(make_para(doc, 'Model form (multiplicative, seasonal period s = 12 months):', bold=True))
elements.append(make_eq(doc, 'Y_t = T_t x S_t x R_t'))
elements.append(make_empty(doc))

elements.append(make_para(doc, 'Trend component -- fitted by OLS linear regression on the extracted trend values (R^2 = 0.9887):', bold=True))
elements.append(make_eq(doc, 'T_hat_t = 100.6499 + 2.6669 * t'))
elements.append(make_para(doc, 'where t = 1 corresponds to the first non-missing trend value (July 1949). After extrapolation, the trend at December 1960 (t = 138) is T_hat = 468.7.'))
elements.append(make_empty(doc))

elements.append(make_para(doc, 'Seasonal factors S_m -- monthly multiplicative indices (product = 1, mean approximately 1.0):', bold=True))
elements.append(make_eq(doc, 'Jan = 0.9102    Feb = 0.8836    Mar = 1.0074    Apr = 0.9759'))
elements.append(make_eq(doc, 'May = 0.9814    Jun = 1.1128    Jul = 1.2266    Aug = 1.2199'))
elements.append(make_eq(doc, 'Sep = 1.0605    Oct = 0.9218    Nov = 0.8012    Dec = 0.8988'))
elements.append(make_empty(doc))

elements.append(make_para(doc, 'The seasonal factors range from 0.8012 (November trough) to 1.2266 (July peak). A value of S_Jul = 1.2266 means July passenger counts are, on average, 22.66% above the trend level. Conversely, S_Nov = 0.8012 means November counts are 19.88% below trend.', bold=False))
elements.append(make_empty(doc))

elements.append(make_para(doc, 'h-step-ahead forecast (multiplicative combination):', bold=True))
elements.append(make_eq(doc, 'Y_hat_{t+h} = (100.6499 + 2.6669 * (t+h)) * S_{month(t+h)}'))
elements.append(make_empty(doc))

elements.append(make_para(doc, '95% prediction interval (based on empirical log-error distribution, log-error SD = 0.0333):', bold=True))
elements.append(make_eq(doc, 'PI_95:  Y_hat_{t+h} * exp(-/+ 1.96 * 0.0333)'))
elements.append(make_para(doc, 'The mean error ratio is 0.9982 (ideal = 1.0 for multiplicative decomposition), confirming the model is well-centred.'))
elements.append(make_empty(doc))

elements.append(make_para(doc, 'Example -- Forecast for July 1961 (h = 7 months after December 1960):', bold=True))
elements.append(make_eq(doc, 'T_hat(Jul 1961) = 100.6499 + 2.6669 * 145  =  487.4'))
elements.append(make_eq(doc, 'Y_hat(Jul 1961) = 487.4 * 1.2266  =  597.8 thousand passengers'))
elements.append(make_eq(doc, '95% PI = [597.8 * exp(-0.0653), 597.8 * exp(+0.0653)]  =  [560.0, 638.1]'))
elements.append(make_empty(doc))

elements.append(make_para(doc, 'Model quality on historical data (144 months):', bold=True))
elements.append(make_bullet(doc, 'R^2 = 0.9918 -- the multiplicative decomposition explains 99.18% of the variance in the observed series.'))
elements.append(make_bullet(doc, 'MAPE = 2.44% -- average absolute percentage error under 3%, indicating excellent fit.'))
elements.append(make_bullet(doc, 'RMSE = 9.88 thousand passengers; MAE = 6.60 thousand passengers.'))

# === Insert all elements before Section 3 ===
for el in reversed(elements):
    body.insert(target_list_idx, el)

doc.save(r'C:\Users\swj17\.claude\projects\R2\Practical_Work_2_Report.docx')
print('Successfully added Section 2.2 with specific fitted model expressions!')
print(f'Added {len(elements)} elements (headings, equations, paragraphs, bullets).')
